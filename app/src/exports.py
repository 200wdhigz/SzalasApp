import pandas as pd
from io import BytesIO
from docx import Document
from reportlab.lib.pagesizes import letter, landscape, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.units import mm
import qrcode
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from flask import send_file
import os

# Rejestracja czcionki obsługującej polskie znaki (jeśli dostępna)
# Szukamy najpierw bundlowanej czcionki, potem systemowej.
PDF_FONT = 'Helvetica'
PDF_FONT_BOLD = 'Helvetica-Bold'


def _try_register_font(font_name: str, font_path: str, bold_name: str = None, bold_path: str = None) -> bool:
    try:
        if os.path.exists(font_path):
            pdfmetrics.registerFont(TTFont(font_name, font_path))
        else:
            return False

        # Bold jest opcjonalny – ReportLab potrzebuje osobnej rejestracji.
        if bold_name and bold_path and os.path.exists(bold_path):
            pdfmetrics.registerFont(TTFont(bold_name, bold_path))
        return True
    except Exception:
        return False


# 1) Repo (najlepsze – deterministyczne), jeśli ktoś doda ./assets/fonts
_repo_dir = os.path.dirname(__file__)
_assets_fonts = os.path.join(_repo_dir, '..', 'static', 'assets', 'fonts')
# typowe nazwy
_dejavu = os.path.join(_assets_fonts, 'DejaVuSans.ttf')
_dejavu_bold = os.path.join(_assets_fonts, 'DejaVuSans-Bold.ttf')

if _try_register_font('DejaVuSans', _dejavu, 'DejaVuSans-Bold', _dejavu_bold):
    PDF_FONT = 'DejaVuSans'
    PDF_FONT_BOLD = 'DejaVuSans-Bold'
else:
    # 2) Windows: Arial (zwykle jest na serwerze Windows) – obsługuje polskie znaki
    _win_fonts = os.path.join(os.environ.get('WINDIR', 'C:\\Windows'), 'Fonts')
    _arial = os.path.join(_win_fonts, 'arial.ttf')
    _arial_bold = os.path.join(_win_fonts, 'arialbd.ttf')
    if _try_register_font('ArialUnicode', _arial, 'ArialUnicode-Bold', _arial_bold):
        PDF_FONT = 'ArialUnicode'
        PDF_FONT_BOLD = 'ArialUnicode-Bold'


def apply_zhp_template_docx(document, title):
    """Nakłada uproszczony szablon ZHP (Nagłówek/Stopka) na dokument Word."""
    section = document.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "ZWIĄZEK HARCERSTWA POLSKIEGO\nSZCZEP SZALAS"
    header_para.style.font.bold = True

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Generowano: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')} | Katalog Sprzętu SzałasApp"
    footer_para.alignment = 1 # Center

def export_to_csv(data, filename):
    df = pd.DataFrame(data)
    output = BytesIO()
    df.to_csv(output, index=False, encoding='utf-8-sig')
    output.seek(0)
    return send_file(output, mimetype='text/csv', as_attachment=True, download_name=f"{filename}.csv")

def export_to_xlsx(data, filename, columns=None):
    df = pd.DataFrame(data)
    if columns:
        # Dodaj brakujące kolumny z pustymi wartościami
        for col in columns:
            if col not in df.columns:
                df[col] = ""
        df = df[columns]

    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False)
    output.seek(0)
    return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', as_attachment=True, download_name=f"{filename}.xlsx")

def export_to_docx(data, filename, title, columns=None):
    document = Document()
    apply_zhp_template_docx(document, title)
    document.add_heading(title, 0)
    
    if not data:
        document.add_paragraph("Brak danych do wyświetlenia.")
    else:
        if columns:
            keys = columns
        else:
            keys = list(data[0].keys())

        table = document.add_table(rows=1, cols=len(keys))
        hdr_cells = table.rows[0].cells
        for i, key in enumerate(keys):
            hdr_cells[i].text = str(key).capitalize()
        
        for item in data:
            row_cells = table.add_row().cells
            for i, key in enumerate(keys):
                row_cells[i].text = str(item.get(key, ''))
                
    output = BytesIO()
    document.save(output)
    output.seek(0)
    return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, download_name=f"{filename}.docx")

def export_to_pdf(data, filename, title, columns=None):
    output = BytesIO()
    doc = SimpleDocTemplate(output, pagesize=landscape(letter))
    elements = []

    styles = getSampleStyleSheet()

    # Ustaw font w stylach bazowych, żeby Paragraph obsługiwał polskie znaki
    try:
        styles['Normal'].fontName = PDF_FONT
        styles['Title'].fontName = PDF_FONT_BOLD if PDF_FONT_BOLD != 'Helvetica-Bold' else PDF_FONT
    except Exception:
        # Jeśli nie uda się ustawić niestandardowej czcionki, użyj domyślnej bez przerywania generowania PDF.
        pass

    # Nagłówek ZHP
    elements.append(Paragraph("<b>ZWIĄZEK HARCERSTWA POLSKIEGO</b>", styles['Normal']))
    elements.append(Paragraph("<b>SZCZEP SZALAS</b>", styles['Normal']))
    elements.append(Spacer(1, 12))

    elements.append(Paragraph(title, styles['Title']))
    elements.append(Spacer(1, 12))

    if not data:
        elements.append(Paragraph("Brak danych.", styles['Normal']))
    else:
        if columns:
            keys = columns
        else:
            keys = list(data[0].keys())

        table_data = [[str(k).capitalize() for k in keys]]
        for item in data:
            row = []
            for k in keys:
                val = str(item.get(k, ''))
                row.append(val)
            table_data.append(row)

        t = Table(table_data)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            # Font dla całej tabeli (i nagłówka i danych) – kluczowe dla polskich znaków
            ('FONTNAME', (0, 0), (-1, 0), PDF_FONT_BOLD),
            ('FONTNAME', (0, 1), (-1, -1), PDF_FONT),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        elements.append(t)

    # Stopka
    elements.append(Spacer(1, 24))
    elements.append(Paragraph(f"Generowano: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}", styles['Normal']))

    doc.build(elements)
    output.seek(0)
    return send_file(output, mimetype='application/pdf', as_attachment=True, download_name=f"{filename}.pdf")

def export_qr_codes_pdf(data, filename, base_url):
    """
    Generuje PDF z kodami QR dla listy przedmiotów.
    Każda etykieta zawiera ID oraz duży kod QR.
    Dodaje linie cięcia i zachowuje safe space.
    """
    output = BytesIO()
    # Używamy A4 w pionie
    doc = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=10 * mm,
        leftMargin=10 * mm,
        topMargin=10 * mm,
        bottomMargin=10 * mm
    )
    elements = []
    styles = getSampleStyleSheet()

    # Ustaw fonty dla polskich znaków
    try:
        styles['Normal'].fontName = PDF_FONT
    except Exception:
        # Jeśli nie uda się ustawić niestandardowej czcionki, użyj domyślnej bez przerywania generowania PDF.
        pass

    if not data:
        elements.append(Paragraph("Brak danych do wygenerowania kodów QR.", styles['Normal']))
    else:
        # Parametry siatki (np. 3 kolumny, 4 rzędy na stronę)
        cols = 3
        rows_per_page = 4
        items_per_page = cols * rows_per_page

        # Create a new style to avoid modifying the shared style object
        id_style = ParagraphStyle(
            'IDStyle',
            parent=styles['Normal'],
            alignment=1,  # Center
            fontSize=10,
            fontName=PDF_FONT_BOLD
        )

        def chunk_data(data, size):
            for i in range(0, len(data), size):
                yield data[i:i + size]

        for page_items in chunk_data(data, items_per_page):
            table_data = []
            current_row = []

            for item in page_items:
                item_id = item.get('id', 'N/A')
                qr_data = f"{base_url}/sprzet/{item_id}"

                # Generowanie QR do BytesIO
                qr = qrcode.QRCode(version=1, box_size=10, border=2) # border=2 to quiet space
                qr.add_data(qr_data)
                qr.make(fit=True)
                qr_img = qr.make_image(fill_color="black", back_color="white")

                img_buffer = BytesIO()
                qr_img.save(img_buffer, format='PNG')
                img_buffer.seek(0)

                # Tworzenie komórki etykiety
                img = Image(img_buffer, width=40 * mm, height=40 * mm)

                # Zawartość etykiety: ID na górze, potem QR
                label_content = [
                    Paragraph(f"<b>{item_id}</b>", id_style),
                    Spacer(1, 1 * mm),
                    img
                ]

                current_row.append(label_content)

                if len(current_row) == cols:
                    table_data.append(current_row)
                    current_row = []

            if current_row:
                # Dopełnij ostatni rząd pustymi komórkami
                while len(current_row) < cols:
                    current_row.append("")
                table_data.append(current_row)

            if table_data:
                # Tabela z liniami cięcia (grid)
                t = Table(table_data, colWidths=[60 * mm] * cols, rowHeights=[55 * mm] * len(table_data))
                t.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.lightgrey), # Linie cięcia
                    ('TOPPADDING', (0, 0), (-1, -1), 2 * mm),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 2 * mm),
                ]))
                elements.append(t)
                from reportlab.platypus import PageBreak
                elements.append(PageBreak())

    doc.build(elements)
    output.seek(0)
    return send_file(output, mimetype='application/pdf', as_attachment=True, download_name=f"{filename}_QR.pdf")
